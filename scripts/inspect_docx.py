# -*- coding: utf-8 -*-
from docx import Document
import sys

path = sys.argv[1]
out = sys.argv[2] if len(sys.argv) > 2 else path + ".inspect.txt"
doc = Document(path)

lines = []
lines.append("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        lines.append(f"P{i}: {t[:200]}")

lines.append("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    lines.append(f"\nTABLE {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows):
        cells = [c.text.replace("\n", " | ")[:100] for c in row.cells]
        lines.append(f"  R{ri}: {cells}")

with open(out, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))
print(out)
