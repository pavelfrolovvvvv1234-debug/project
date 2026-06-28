# -*- coding: utf-8 -*-
"""Fill PM02 and PM11 practice report docx files."""
import shutil
import tempfile
from pathlib import Path

from docx import Document


def set_para(p, text: str) -> None:
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def set_cell(cell, text: str) -> None:
    cell.text = text


GITHUB = "https://github.com/pavelfrolovvvvv1234-debug/project"
COMMITS = "https://github.com/pavelfrolovvvvv1234-debug/project/commits/main"

PM02_REPORT = [
    (
        "В ходе практики разработана веб-система SiteNet Manager — панель для управления сетью сайтов. "
        "Проект состоит из клиентской части (React + Vite + TypeScript) и серверной (Express + TypeScript). "
        "Модули связаны через REST API: фронтенд отправляет запросы с JWT-токеном, бэкенд обрабатывает их "
        "и работает с PostgreSQL через Prisma ORM. Данные передаются в формате JSON. Документация API "
        "оформлена в Swagger (/api/docs). Клиент и сервер запускаются отдельно: фронт на порту 5173, "
        "API на 3001. При деплое фронт размещается на Vercel, бэкенд и БД — на Render."
    ),
    (
        "По заданию система должна вести учёт сетей сайтов, проверять их доступность, собирать данные "
        "по страницам и управлять перелинковкой. Реализованы: регистрация и вход (JWT), CRUD сетей и сайтов, "
        "дашборд со статистикой, массовая проверка HTTP-статусов, проверка страницы сайта (title, description, "
        "h1, robots.txt, Open Graph, sitemap), uptime-проверки, матрица ссылок между сайтами в сети, экспорт "
        "CSV и sitemap.xml, облако ключевых слов, поведенческие метрики. Каждый пользователь видит только "
        "свои сети и сайты — доступ проверяется на уровне API."
    ),
    (
        "Модули разбиты по зонам ответственности. На бэкенде: auth, networks, sites, links, keywords, "
        "dashboard, templates. На фронтенде: страницы Login, Dashboard, Networks, Sites, детальная страница "
        "сети (вкладки: сайты, матрица, ключи, экспорт) и сайта (сводка, проверки, ссылки, ключи). "
        "Сервис seo.ts парсит HTML через cheerio, behavioral.ts считает метрики вовлечённости. "
        "Схема БД: User, Network, Site, SiteLink, SeoSnapshot, UptimeCheck, SiteKeyword, BehavioralSnapshot, "
        "Tag, SiteTemplate. Связи между таблицами настроены через внешние ключи в Prisma."
    ),
    (
        "Тестирование проводилось вручную через интерфейс и Swagger. Проверены: вход под demo-аккаунтом, "
        "создание/редактирование/удаление сети и сайта, SEO-проверка на реальных доменах (github.com, "
        "wikipedia.org), uptime-проверка, создание и удаление ссылок в матрице, экспорт CSV и sitemap, "
        "массовая проверка всех сайтов с дашборда. Ошибки валидации (пустые поля, неверный токен) "
        "возвращают коды 400/401/404. На seed-данных с фейковыми доменами проверки показывают ошибку — "
        "это ожидаемое поведение."
    ),
    (
        "Исходный код выложен на GitHub с историей из 20+ коммитов. В репозитории есть README с инструкцией "
        "запуска, docker-compose.yml для локальной PostgreSQL, render.yaml для деплоя бэкенда, docs/schema.sql "
        "и docs/postgresql-extras.sql (представления, функции, триггеры). При изменениях схемы применялись "
        "миграции Prisma. API описан в Swagger, что упрощает дальнейшую доработку без разбора всего кода."
    ),
]

PM02_DIARY = [
    "Изучил задание, настроил Node, Git, структуру monorepo",
    "Поднял Express, Prisma, схему БД, JWT-авторизацию",
    "CRUD сетей и сайтов, middleware проверки владельца",
    "SEO и uptime-проверки, парсинг HTML",
    "Матрица ссылок, валидация, экспорт CSV",
    "React-фронт: роутинг, login, layout",
    "Страницы Dashboard, Networks, Sites",
    "Детальная страница сайта, SEO-панель",
    "Облако ключевых слов, поведенческие метрики",
    "Дашборд: массовая HTTP-проверка",
    "Swagger, seed-данные, тестирование API",
    "Расширенный аудит страницы (OG, robots, sitemap)",
    "Деплой на Render/Vercel, правка CORS и env",
    "Финальное тестирование, правки по замечаниям",
    "Оформление отчёта, подготовка к сдаче",
]

PM02_CHAR = (
    "За время практики освоил разработку fullstack-приложения на TypeScript: клиент на React, "
    "сервер на Express, БД PostgreSQL через Prisma. Научился проектировать REST API, настраивать "
    "JWT-аутентификацию, связывать фронтенд и бэкенд через HTTP-запросы. Реализовал CRUD-операции, "
    "бизнес-логику проверки сайтов, экспорт данных. Применил миграции БД, seed-данные, Docker для "
    "локальной разработки. Разместил проект на GitHub, оформил README и Swagger-документацию. "
    "Может самостоятельно дорабатывать и сопровождать подобные веб-системы."
)

PM11_DIARY = [
    "Проектирование ERD, выбор PostgreSQL",
    "Prisma schema: users, networks, sites",
    "Таблицы seo_snapshots, uptime_checks, site_links",
    "site_keywords, behavioral_snapshots, tags",
    "Миграции, индексы, unique-ограничения",
    "seed.ts — тестовые данные",
    "schema.sql — дамп структуры",
    "postgresql-extras.sql: views, functions, triggers",
    "Проверка запросов через API и Prisma Studio",
    "Деплой БД на Render, migrate deploy",
]

PM11_CHAR = (
    "Освоил проектирование реляционных БД: ERD, нормализация, внешние ключи, индексы. "
    "Работал с PostgreSQL и Prisma ORM: миграции, seed, CRUD через клиент. Написал SQL-скрипты "
    "для представлений, функций и триггеров. Понял, как БД связана с REST API и фронтендом. "
    "Может самостоятельно проектировать и сопровождать БД для веб-приложений."
)

PM11_DESC = (
    "SiteNet Manager — веб-панель для управления сетью сайтов. Пользователь — владелец или SEO-специалист, "
    "которому нужно вести учёт доменов, проверять их доступность и SEO-параметры, строить перелинковку "
    "между сайтами в одной сети и анализировать ключевые слова."
)

PM11_FEATURES = """— PostgreSQL-схема с 11 таблицами и внешними ключами
— Миграции Prisma и seed-данные
— SQL-скрипты schema.sql и postgresql-extras.sql (views, functions, triggers)
— JWT-авторизация с привязкой данных к пользователю
— CRUD сетей и сайтов через API
— SEO-снимки и uptime-проверки в БД
— Матрица перелинковки (site_links)
— Ключевые слова и поведенческие метрики
— Экспорт sitemap.xml и CSV"""

PM11_ARCH = (
    "Схема: UI (React) → REST API (Express) → PostgreSQL (Prisma). Пользователь работает в браузере. "
    "React-приложение шлёт запросы на /api/* с Bearer-токеном. Express-роуты вызывают Prisma Client, "
    "который выполняет SQL к PostgreSQL. Ответ — JSON. При деплое фронт на Vercel, API и БД на Render."
)

PM11_ENTITIES = """users — пользователи (email, пароль, роль)
networks — сети сайтов (название, описание, владелец)
sites — сайты в сети (домен, title, ниша, статус)
site_links — ссылки между сайтами (from, to, anchor, тип)
seo_snapshots — результаты проверки страницы
uptime_checks — результаты проверки доступности
site_keywords — ключевые слова сайта
behavioral_snapshots — поведенческие метрики
tags, site_tags — теги сайтов
site_templates — HTML-шаблоны лендингов

Связи: User → Networks → Sites; Site → SeoSnapshots, UptimeChecks, Keywords; SiteLink связывает два Site в одной Network."""

PM11_USECASES = """UC1 — Вход в систему
Пользователь вводит email и пароль → API проверяет → возвращает JWT → фронт сохраняет токен → открывается дашборд.

UC2 — Добавить сайт в сеть
Пользователь выбирает сеть → заполняет домен и название → POST /api/sites → запись в sites → сайт в списке.

UC3 — Проверить страницу сайта
Пользователь открывает сайт → «Проверить страницу» → POST /api/sites/:id/seo-check → парсинг HTML → запись в seo_snapshots → сводка на экране.

UC4 — Матрица перелинковки
Пользователь на вкладке «Матрица» → кликает ячейку → запись в site_links → матрица обновляется."""

PM11_API = [
    ("POST", "/api/auth/register", "Регистрация", "email, password → 201"),
    ("POST", "/api/auth/login", "Вход, JWT", "email, password → token"),
    ("GET", "/api/auth/me", "Текущий пользователь", "Bearer → user"),
    ("GET", "/api/dashboard", "Статистика", "→ counts, avg score"),
    ("POST", "/api/dashboard/check-all", "Проверить все сайты", "→ results[]"),
    ("GET/POST", "/api/networks", "Список / создание", "name, description"),
    ("GET/PUT/DELETE", "/api/networks/:id", "Сеть по id", "CRUD"),
    ("GET", "/api/networks/:id/sitemap.xml", "Sitemap сети", "→ XML"),
    ("GET", "/api/networks/:id/export-links.csv", "Экспорт ссылок", "→ CSV"),
    ("GET/POST", "/api/sites", "Список / создание", "domain, title"),
    ("GET/PUT/DELETE", "/api/sites/:id", "Сайт по id", "CRUD"),
    ("POST", "/api/sites/:id/seo-check", "Проверка страницы", "→ snapshot"),
    ("POST", "/api/sites/:id/uptime-check", "Uptime", "→ check"),
    ("POST", "/api/sites/:id/behavioral-check", "Поведенческие метрики", "→ snapshot"),
    ("GET/POST", "/api/links", "Ссылки", "from, to, anchor"),
    ("GET", "/api/links/network/:id/matrix", "Матрица", "→ matrix"),
    ("GET", "/api/keywords/network/:id", "Ключи сети", "→ keywords[]"),
]

PM11_TRACE = [
    ("Схема БД", f"{GITHUB}/tree/main/backend/prisma", "[URL деплоя]"),
    ("Авторизация", f"{GITHUB}/blob/main/backend/src/routes/auth.ts", "[URL]/login"),
    ("CRUD сетей", f"{GITHUB}/blob/main/backend/src/routes/networks.ts", "[URL]/networks"),
    ("CRUD сайтов", f"{GITHUB}/blob/main/backend/src/routes/sites.ts", "[URL]/sites"),
    ("Матрица ссылок", f"{GITHUB}/blob/main/backend/src/routes/links.ts", "[URL]/networks/:id"),
    ("Фронтенд", f"{GITHUB}/tree/main/frontend/src", "[URL]"),
    ("Дашборд", f"{GITHUB}/blob/main/frontend/src/pages/DashboardPage.tsx", "[URL]/"),
    ("SEO-аудит", f"{GITHUB}/blob/main/backend/src/services/seo.ts", "[URL]/sites/:id"),
    ("SQL extras", f"{GITHUB}/blob/main/docs/postgresql-extras.sql", "—"),
]

PM11_CONCLUSION = (
    "В ходе практики по ПМ11 спроектирована и реализована реляционная БД PostgreSQL для SiteNet Manager. "
    "Созданы таблицы с внешними ключами и индексами, миграции Prisma, seed-данные, дополнительные SQL-объекты "
    "(представления, функции, триггеры). БД интегрирована с REST API: все операции в интерфейсе сохраняются "
    "и читаются из PostgreSQL. Проект размещён на GitHub. Сложнее всего было настроить связи между таблицами "
    "и миграции при смене СУБД с SQLite на PostgreSQL. Прокачал навыки проектирования ERD, работы с Prisma "
    "и администрирования БД на Render. В дальнейшем можно добавить бэкапы и роли на уровне PostgreSQL."
)


def save_doc(doc: Document, path: str) -> None:
    target = Path(path)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = tmp.name
    doc.save(tmp_path)
    shutil.move(tmp_path, target)


def fill_pm02(path_in: str, path_out: str) -> None:
    doc = Document(path_in)

    placeholders = [129, 133, 137, 141, 145]
    for idx, pnum in enumerate(placeholders):
        set_para(doc.paragraphs[pnum], PM02_REPORT[idx])

    diary = doc.tables[4]
    hours = ["8"] * 14 + ["4"]
    for i, (text, h) in enumerate(zip(PM02_DIARY, hours), start=1):
        if i < len(diary.rows):
            set_cell(diary.rows[i].cells[1], text)
            set_cell(diary.rows[i].cells[3], h)

    set_para(doc.paragraphs[105], PM02_CHAR)

    save_doc(doc, path_out)


def fill_pm11(path_in: str, path_out: str) -> None:
    doc = Document(path_in)

    set_para(doc.paragraphs[130], "Название проекта: SiteNet Manager")
    set_para(doc.paragraphs[131], f"Ссылка на проект из каталога: {GITHUB}")
    set_para(doc.paragraphs[133], PM11_DESC)
    set_para(doc.paragraphs[135], PM11_FEATURES)
    set_para(doc.paragraphs[143], PM11_ARCH)
    set_para(doc.paragraphs[148], PM11_ENTITIES)
    set_para(doc.paragraphs[153], PM11_USECASES)
    set_para(doc.paragraphs[164], PM11_CONCLUSION)

    set_para(doc.paragraphs[106], "Работал (а) с 15 июня по 28 июня 2026 года и показал(а) результаты по освоению следующих общих компетенций")

    passport = doc.tables[10]
    set_cell(passport.rows[0].cells[1], "SiteNet Manager")
    set_cell(passport.rows[1].cells[1], GITHUB)
    set_cell(passport.rows[2].cells[1], GITHUB)
    set_cell(passport.rows[3].cells[1], "[URL Vercel] — фронт, [URL Render] — API")
    set_cell(passport.rows[4].cells[1], "React, Vite, TypeScript, Tailwind")
    set_cell(passport.rows[5].cells[1], "Node.js, Express, TypeScript")
    set_cell(passport.rows[6].cells[1], "PostgreSQL")
    set_cell(passport.rows[7].cells[1], "Да")
    set_cell(passport.rows[8].cells[1], "login: demo@college.local   password: Demo123!")

    usecase_table = doc.tables[11]
    scenarios = [
        ("Вход в систему", "1. Открыть /login 2. Ввести email и пароль 3. Нажать «Войти» 4. Открывается дашборд"),
        ("Добавить сайт", "1. Открыть сеть 2. Вкладка «Сайты» 3. Заполнить домен и название 4. Сохранить"),
        ("Проверить страницу", "1. Открыть сайт 2. «Проверить страницу» 3. Дождаться ответа 4. Смотреть сводку"),
        ("Матрица ссылок", "1. Открыть сеть 2. Вкладка «Матрица» 3. Клик по ячейке 4. Ссылка создана"),
    ]
    for i, (name, steps) in enumerate(scenarios, start=1):
        set_cell(usecase_table.rows[i].cells[0], name)
        set_cell(usecase_table.rows[i].cells[1], steps)

    api_table = doc.tables[12]
    for i, row in enumerate(PM11_API[:5], start=1):
        method, url, purpose, example = row
        set_cell(api_table.rows[i].cells[0], method)
        set_cell(api_table.rows[i].cells[1], url)
        set_cell(api_table.rows[i].cells[2], purpose)
        set_cell(api_table.rows[i].cells[3], example)

    trace = doc.tables[13]
    for i, (uc, gh, deploy) in enumerate(PM11_TRACE, start=1):
        set_cell(trace.rows[i].cells[0], uc)
        set_cell(trace.rows[i].cells[1], gh)
        set_cell(trace.rows[i].cells[2], deploy)

    demo = doc.tables[14]
    set_cell(demo.rows[0].cells[1], "GIF")
    set_cell(demo.rows[1].cells[1], "[URL на GIF — вставить после записи]")
    set_cell(demo.rows[2].cells[1], "Вход, дашборд, сеть, матрица, проверка страницы сайта")
    set_cell(demo.rows[3].cells[1], "Демо-аккаунт: demo@college.local / Demo123!")

    quality = doc.tables[15]
    set_cell(quality.rows[0].cells[1], "Да")
    set_cell(quality.rows[1].cells[1], "[Да — после подключения Code Climate]")
    set_cell(quality.rows[2].cells[1], "[A или B]")
    set_cell(quality.rows[3].cells[1], "Да")
    set_cell(quality.rows[4].cells[1], f"20+ коммитов, история: {COMMITS}")

    diary = doc.tables[4]
    hours = ["8", "8", "8", "8", "8", "8", "8", "8", "8", "6"]
    for i, (text, h) in enumerate(zip(PM11_DIARY, hours), start=1):
        if i < len(diary.rows):
            set_cell(diary.rows[i].cells[1], text)
            set_cell(diary.rows[i].cells[3], h)

    set_para(doc.paragraphs[105], PM11_CHAR)

    save_doc(doc, path_out)


if __name__ == "__main__":
    pm02 = r"c:\Users\xd-user\Downloads\Отчетная документация по ПП ПМ02.docx"
    pm11 = r"c:\Users\xd-user\Downloads\Отчетная документация по ПП ПМ11 (2).docx"

    fill_pm02(pm02, pm02)
    fill_pm11(pm11, pm11)
